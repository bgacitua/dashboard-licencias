import io
import os
from typing import List, Optional
from decimal import Decimal
from sqlalchemy.orm import Session

from app.repositories.seleccion_repository import SeleccionRepository
from app.schemas.seleccion import CandidatoCreate, CandidatoUpdate
from app.models.seleccion import CandidatoSeleccion
from app.core.logging_config import logger


TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "carta_oferta.docx")


class SeleccionService:
    def __init__(self, db: Session):
        self.repo = SeleccionRepository(db)

    def get_all(self):
        return self.repo.get_all()

    def get_by_id(self, candidato_id: int):
        return self.repo.get_by_id(candidato_id)

    def get_by_rut(self, rut: str):
        return self.repo.get_by_rut(rut)

    def create(self, data: CandidatoCreate):
        return self.repo.create(data)

    def update(self, candidato_id: int, data: CandidatoUpdate):
        return self.repo.update(candidato_id, data)

    def delete(self, candidato_id: int) -> bool:
        return self.repo.delete(candidato_id)

    def generar_carta_oferta(self, candidato_id: int) -> bytes:
        from docx import Document
        from docx.shared import Pt

        candidato = self.repo.get_by_id(candidato_id)
        if not candidato:
            raise ValueError(f"Candidato {candidato_id} no encontrado")

        if os.path.exists(TEMPLATE_PATH):
            doc = Document(TEMPLATE_PATH)
            _reemplazar_marcadores(doc, candidato)
        else:
            doc = _generar_carta_desde_cero(candidato)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


def _fmt_monto(valor) -> str:
    if valor is None:
        return "$0"
    try:
        return f"${int(valor):,}".replace(",", ".")
    except Exception:
        return str(valor)


def _fmt_fecha(fecha) -> str:
    if fecha is None:
        return ""
    try:
        return fecha.strftime("%d de %B de %Y")
    except Exception:
        return str(fecha)


MARCADORES = {
    "{{empresa}}": lambda c: c.empresa or "",
    "{{nombre}}": lambda c: c.nombre or "",
    "{{rut}}": lambda c: c.rut or "",
    "{{cargo}}": lambda c: c.cargo or "",
    "{{lugar_de_trabajo}}": lambda c: c.lugar_de_trabajo or "",
    "{{jornada_de_trabajo}}": lambda c: c.jornada_de_trabajo or "",
    "{{tipo_de_contrato}}": lambda c: c.tipo_de_contrato or "",
    "{{fecha_de_inicio}}": lambda c: _fmt_fecha(c.fecha_de_inicio),
    "{{gerencia}}": lambda c: c.gerencia or "",
    "{{sueldo_base}}": lambda c: _fmt_monto(c.sueldo_base),
    "{{bono}}": lambda c: _fmt_monto(c.bono),
    "{{movilizacion}}": lambda c: _fmt_monto(c.movilizacion),
    "{{correo_analista}}": lambda c: c.correo_analista or "",
    "{{fecha_cierre}}": lambda c: _fmt_fecha(c.fecha_cierre),
}


def _reemplazar_marcadores(doc, candidato):
    for para in doc.paragraphs:
        for run in para.runs:
            for marcador, fn in MARCADORES.items():
                if marcador in run.text:
                    run.text = run.text.replace(marcador, fn(candidato))
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for para in celda.paragraphs:
                    for run in para.runs:
                        for marcador, fn in MARCADORES.items():
                            if marcador in run.text:
                                run.text = run.text.replace(marcador, fn(candidato))


def _generar_carta_desde_cero(candidato):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("CARTA DE OFERTA")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    fecha_hoy = _fmt_fecha(__import__("datetime").date.today())
    doc.add_paragraph(f"{candidato.empresa or ''}, {fecha_hoy}")
    doc.add_paragraph()

    doc.add_paragraph(f"Estimado/a {candidato.nombre or ''},")
    doc.add_paragraph()

    intro = doc.add_paragraph(
        f"Es un agrado para nosotros hacerle llegar la presente carta de oferta laboral, "
        f"la que detalla las condiciones bajo las cuales {candidato.empresa or 'la empresa'} "
        f"le propone integrarse a nuestro equipo de trabajo."
    )

    doc.add_paragraph()

    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"

    def agregar_fila(clave, valor):
        fila = tabla.add_row().cells
        fila[0].text = clave
        fila[1].text = valor

    agregar_fila("Nombre", candidato.nombre or "")
    agregar_fila("RUT", candidato.rut or "")
    agregar_fila("Cargo", candidato.cargo or "")
    agregar_fila("Gerencia", candidato.gerencia or "")
    agregar_fila("Lugar de Trabajo", candidato.lugar_de_trabajo or "")
    agregar_fila("Jornada", candidato.jornada_de_trabajo or "")
    agregar_fila("Tipo de Contrato", candidato.tipo_de_contrato or "")
    agregar_fila("Fecha de Inicio", _fmt_fecha(candidato.fecha_de_inicio))
    agregar_fila("Sueldo Base", _fmt_monto(candidato.sueldo_base))
    agregar_fila("Bono", _fmt_monto(candidato.bono))
    agregar_fila("Movilización", _fmt_monto(candidato.movilizacion))

    doc.add_paragraph()
    doc.add_paragraph(
        "Esta oferta está sujeta a la verificación de antecedentes y firma de contrato. "
        "Le rogamos confirmar su aceptación a la brevedad posible."
    )
    doc.add_paragraph()
    doc.add_paragraph("Atentamente,")
    doc.add_paragraph()
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Analista de Selección")
    if candidato.correo_analista:
        doc.add_paragraph(candidato.correo_analista)

    return doc
